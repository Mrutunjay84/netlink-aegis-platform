"""Netlink Aegis - company document template for policy export.

Lets an admin define a reusable "letterhead" applied at export time so every
drafted policy comes out branded:

  - a **company name** + **header/footer text**, and a **watermark** (text), or
  - a full **uploaded .docx template** (the admin designs the header / footer /
    watermark / logo in Word; generated content is appended into it for DOCX
    export).

Settings live in the existing ``GlobalSettings(name="general")`` row under
``netlink_doc_*`` keys (the uploaded .docx is stored base64-encoded there, so it
persists with the database and needs no media volume). The per-export toggle is
sent by the client; this module just provides the building blocks.
"""

from __future__ import annotations

import base64
import io

import structlog

logger = structlog.get_logger(__name__)


# Field keys (in GlobalSettings "general").
ENABLED = "netlink_doc_template_enabled"
COMPANY = "netlink_doc_company_name"
WATERMARK = "netlink_doc_watermark_text"
HEADER = "netlink_doc_header_text"
FOOTER = "netlink_doc_footer_text"
DOCX_B64 = "netlink_doc_template_docx_b64"
DOCX_NAME = "netlink_doc_template_docx_name"

TEXT_FIELDS = [COMPANY, WATERMARK, HEADER, FOOTER]
MANAGED_FIELDS = {ENABLED, COMPANY, WATERMARK, HEADER, FOOTER, DOCX_B64, DOCX_NAME}


# ---------------------------------------------------------------------------
# Settings storage (reuses the general GlobalSettings row)
# ---------------------------------------------------------------------------
def _general_value() -> dict:
    try:
        from global_settings.models import GlobalSettings

        gs = GlobalSettings.objects.filter(name="general").first()
        if gs and isinstance(gs.value, dict):
            return gs.value
    except Exception as e:
        logger.warning("netlink_doc_template_load_failed", error=str(e))
    return {}


def get_template_settings() -> dict:
    """Raw settings incl. the base64 docx (server-side use only)."""
    value = _general_value()
    out = {f: str(value.get(f, "") or "") for f in MANAGED_FIELDS}
    return out


def masked_template_settings() -> dict:
    """Client-facing view: text fields + flags, never the docx bytes."""
    s = get_template_settings()
    return {
        "enabled": s.get(ENABLED, "") == "true",
        "company_name": s.get(COMPANY, ""),
        "watermark_text": s.get(WATERMARK, ""),
        "header_text": s.get(HEADER, ""),
        "footer_text": s.get(FOOTER, ""),
        "has_docx": bool(s.get(DOCX_B64, "")),
        "docx_name": s.get(DOCX_NAME, ""),
    }


def save_template_settings(updates: dict) -> dict:
    """Merge managed fields into the general settings row; return masked view."""
    from global_settings.models import GlobalSettings

    gs, _ = GlobalSettings.objects.get_or_create(name="general", defaults={"value": {}})
    value = gs.value if isinstance(gs.value, dict) else {}
    for field, raw in updates.items():
        if field not in MANAGED_FIELDS:
            continue
        value[field] = "" if raw is None else str(raw)
    gs.value = value
    gs.save(update_fields=["value"])
    logger.info("netlink_doc_template_saved", fields=sorted(updates.keys()))
    return masked_template_settings()


def validate_docx_b64(b64: str) -> bool:
    """True if ``b64`` decodes to a readable .docx."""
    try:
        from docx import Document

        Document(io.BytesIO(base64.b64decode(b64)))
        return True
    except Exception:
        return False


# ---------------------------------------------------------------------------
# DOCX application
# ---------------------------------------------------------------------------
def ensure_default_styles(document) -> None:
    """Copy Word's built-in styles into ``document`` if they're missing.

    Uploaded letterheads (especially ones designed in Canva/Word) often strip
    unused styles like ``Heading 1`` / ``List Bullet`` / ``Table Grid``.
    htmldocx looks these up *by name* and raises ``KeyError`` if absent, which
    breaks the export. We append any default style not already present (matched
    by styleId, so the template's own custom styles are left untouched).
    """
    import copy

    from docx import Document as _Document
    from docx.oxml.ns import qn

    try:
        default_styles = _Document().styles.element
        dst = document.styles.element
        existing = {
            s.get(qn("w:styleId")) for s in dst.findall(qn("w:style"))
        }
        for style in default_styles.findall(qn("w:style")):
            sid = style.get(qn("w:styleId"))
            if sid and sid not in existing:
                dst.append(copy.deepcopy(style))
                existing.add(sid)
    except Exception as e:
        logger.warning("netlink_doc_ensure_styles_failed", error=str(e))


def open_base_document(settings: dict):
    """Return a python-docx Document from the uploaded template, or None.

    The returned document is guaranteed to carry Word's built-in styles so
    htmldocx can append headings/lists/tables into it.
    """
    b64 = settings.get(DOCX_B64, "")
    if not b64:
        return None
    try:
        from docx import Document

        document = Document(io.BytesIO(base64.b64decode(b64)))
        ensure_default_styles(document)
        return document
    except Exception as e:
        logger.warning("netlink_doc_template_open_failed", error=str(e))
        return None


def _xml_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _add_text_watermark(document, text: str) -> None:
    """Inject a diagonal VML text watermark into every section header."""
    from docx.oxml import parse_xml

    safe = _xml_escape(text)[:80]
    xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:w10="urn:schemas-microsoft-com:office:word">'
        "<w:rPr><w:noProof/></w:rPr>"
        "<w:pict>"
        '<v:shape id="NetlinkWatermark" o:spid="_x0000_s2049" type="#_x0000_t136" '
        'style="position:absolute;margin-left:0;margin-top:0;width:468pt;'
        "height:117pt;rotation:315;z-index:-251654144;"
        "mso-position-horizontal:center;mso-position-horizontal-relative:margin;"
        'mso-position-vertical:center;mso-position-vertical-relative:margin" '
        'o:allowincell="f" fillcolor="#d9d9d9" stroked="f">'
        '<v:textpath style="font-family:&quot;Calibri&quot;;font-size:1pt" '
        f'string="{safe}"/>'
        "</v:shape>"
        "</w:pict>"
        "</w:r>"
    )
    for section in document.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para._p.append(parse_xml(xml))


def apply_docx_branding(document, settings: dict) -> None:
    """Apply field-based header/footer text + watermark to a fresh document.

    Used when no uploaded .docx template is configured (an uploaded template
    already carries its own branding).
    """
    company = settings.get(COMPANY, "").strip()
    header_text = settings.get(HEADER, "").strip() or company
    footer_text = settings.get(FOOTER, "").strip()
    watermark = settings.get(WATERMARK, "").strip() or company

    for section in document.sections:
        if header_text:
            hp = section.header.paragraphs[0]
            hp.text = header_text
        if footer_text:
            fp = section.footer.paragraphs[0]
            fp.text = footer_text

    if watermark:
        try:
            _add_text_watermark(document, watermark)
        except Exception as e:
            logger.warning("netlink_doc_watermark_failed", error=str(e))


# ---------------------------------------------------------------------------
# PDF branding (WeasyPrint)
# ---------------------------------------------------------------------------
def _css_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace('"', '\\"')


def pdf_branding(settings: dict) -> dict:
    """Return ``{css, watermark_html}`` to inject into the PDF document.

    Adds running header/footer (company / header / footer text + page numbers)
    and a repeated diagonal watermark.
    """
    company = settings.get(COMPANY, "").strip()
    header_text = settings.get(HEADER, "").strip() or company
    footer_text = settings.get(FOOTER, "").strip()
    watermark = settings.get(WATERMARK, "").strip() or company

    css_parts = ["@page {"]
    if company:
        css_parts.append(
            f'  @top-left {{ content: "{_css_escape(company)}"; '
            "font-size: 8pt; color: #999; }"
        )
    if header_text and header_text != company:
        css_parts.append(
            f'  @top-right {{ content: "{_css_escape(header_text)}"; '
            "font-size: 8pt; color: #999; }"
        )
    if footer_text:
        css_parts.append(
            f'  @bottom-left {{ content: "{_css_escape(footer_text)}"; '
            "font-size: 8pt; color: #999; }"
        )
    css_parts.append(
        '  @bottom-right { content: "Page " counter(page) " of " counter(pages); '
        "font-size: 8pt; color: #999; }"
    )
    css_parts.append("}")

    css_parts.append(
        ".netlink-watermark { position: fixed; top: 42%; left: 0; right: 0; "
        "text-align: center; transform: rotate(-45deg); transform-origin: center; "
        "font-size: 92px; font-weight: 700; letter-spacing: 6px; "
        "text-transform: uppercase; color: rgba(120,120,120,0.12); z-index: -1; }"
    )

    watermark_html = ""
    if watermark:
        safe = (
            watermark.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        )
        watermark_html = f'<div class="netlink-watermark">{safe}</div>'

    return {"css": "\n".join(css_parts), "watermark_html": watermark_html}
